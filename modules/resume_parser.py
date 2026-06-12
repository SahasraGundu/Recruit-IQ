# modules/resume_parser.py
import re
import io
import pdfplumber
import PyPDF2
from typing import Dict, Any, List
from docx import Document as DocxDocument


SKILLS_DB = [
    # Programming languages
    "python","java","javascript","typescript","c++","c#","go","rust","kotlin","swift",
    "php","ruby","scala","r","matlab","perl","bash","powershell",
    # Web
    "react","angular","vue","nextjs","nodejs","express","django","flask","fastapi",
    "spring","html","css","tailwind","bootstrap","graphql","rest","soap",
    # Data / AI / ML
    "machine learning","deep learning","nlp","computer vision","tensorflow","pytorch",
    "keras","scikit-learn","pandas","numpy","opencv","huggingface","langchain",
    "llm","generative ai","prompt engineering","rag",
    # Cloud & DevOps
    "aws","azure","gcp","docker","kubernetes","terraform","ansible","jenkins","ci/cd",
    "github actions","linux","nginx","apache",
    # Databases
    "mysql","postgresql","mongodb","redis","elasticsearch","cassandra","dynamodb",
    "sqlite","oracle","mssql","firebase","supabase",
    # Mobile
    "android","ios","react native","flutter","xamarin",
    # Tools
    "git","jira","confluence","figma","postman","selenium","pytest","junit",
    # Data tools
    "tableau","powerbi","excel","spark","hadoop","kafka","airflow","dbt","snowflake",
    # Domain skills
    "agile","scrum","microservices","system design","api design","devops","mlops",
]

INDIAN_INSTITUTIONS = {
    "iit": 10, "nit": 8, "iiit": 8, "bits": 8, "iim": 10,
    "vit": 7, "manipal": 7, "srm": 6, "amity": 5,
    "osmania": 6, "jntu": 6, "ou": 6, "vnr": 6, "cbit": 6,
    "hyderabad": 5, "bangalore": 5, "pune": 5,
}

def extract_text_from_pdf(pdf_file) -> str:
    text = ""
    try:
        pdf_bytes = pdf_file.read() if hasattr(pdf_file, "read") else pdf_file
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception:
        pass
    if not text.strip():
        try:
            pdf_bytes = pdf_file.read() if hasattr(pdf_file, "read") else pdf_file
            reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        except Exception:
            pass
    return text

def extract_text_from_docx(docx_file) -> str:
    text = ""
    try:
        docx_bytes = docx_file.read() if hasattr(docx_file, "read") else docx_file
        doc = DocxDocument(io.BytesIO(docx_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception:
        pass
    return text

def extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    SKIP_WORDS = [
        "resume", "cv", "curriculum", "vitae", "profile",
        "summary", "professional", "objective", "overview",
        "contact", "details", "information", "address"
    ]
    for line in lines[:8]:
        words = line.split()
        if (2 <= len(words) <= 3 and
                not any(c.isdigit() for c in line) and
                not any(w.lower() in SKIP_WORDS for w in words)):
            return line.title()
    return "Unknown"

def extract_email(text: str) -> str:
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(pattern, text)
    return match.group(0).lower() if match else ""

def extract_phone(text: str) -> str:
    patterns = [
        r'(?:\+91[\s\-]?)?[6-9]\d{9}',
        r'\d{3}[\s\-]\d{3}[\s\-]\d{4}',
        r'\(\d{3}\)\s*\d{3}[\s\-]\d{4}',
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(0)
    return ""

def extract_skills(text: str) -> List[str]:
    text_lower = text.lower()
    found = []
    for skill in SKILLS_DB:
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            found.append(skill)
    return list(set(found))

def extract_experience_years(text: str) -> float:
    patterns = [
        r'(\d+\.?\d*)\s*\+?\s*years?\s+(?:of\s+)?(?:work\s+)?experience',
        r'experience\s*[:\-]?\s*(\d+\.?\d*)\s*\+?\s*years?',
        r'(\d+\.?\d*)\s*yrs?\s+(?:of\s+)?experience',
        r'total\s+experience[:\s]+(\d+\.?\d*)',
    ]
    for p in patterns:
        m = re.search(p, text.lower())
        if m:
            return float(m.group(1))
    # Count from job dates
    year_ranges = re.findall(r'(20\d{2})\s*[-–]\s*(20\d{2}|present|current)', text.lower())
    total = 0
    for start, end in year_ranges:
        end_year = 2024 if end in ["present", "current"] else int(end)
        total += max(0, end_year - int(start))
    return float(min(total, 30)) if total > 0 else 0.0

def extract_education(text: str) -> List[Dict]:
    education = []
    degrees = ["b.tech","btech","b.e","be","m.tech","mtech","m.e","me","bca","mca",
               "b.sc","bsc","msc","m.sc","mba","b.com","bcom","phd","ph.d"]
    lines = text.split("\n")
    for i, line in enumerate(lines):
        line_lower = line.lower()
        for deg in degrees:
            if deg in line_lower:
                context = " ".join(lines[max(0, i-1):i+3])
                inst_score = 0
                inst_name = ""
                for inst, score in INDIAN_INSTITUTIONS.items():
                    if inst in context.lower():
                        inst_score = score
                        inst_name = inst.upper()
                        break
                years = re.findall(r'20\d{2}', context)
                education.append({
                    "degree": deg.upper(),
                    "institution": inst_name,
                    "institution_score": inst_score,
                    "years": years[:2] if years else [],
                    "context": context.strip()[:200]
                })
                break
    return education

def extract_ctc(text: str, keyword="current") -> float:
    """Extract CTC in LPA (Lakhs Per Annum)."""
    patterns = [
        rf'{keyword}\s+ctc\s*[:\-]?\s*(?:rs\.?|inr|₹)?\s*(\d+\.?\d*)\s*(?:lpa|lakh|l|lakhs)',
        rf'ctc\s*[:\-]?\s*(?:rs\.?|inr|₹)?\s*(\d+\.?\d*)\s*(?:lpa|lakh|l)',
        rf'(\d+\.?\d*)\s*(?:lpa|lakh|lakhs?)\s*(?:per\s+annum)?',
        rf'(?:rs\.?|inr|₹)\s*(\d+[\d,]*)\s*(?:per\s+annum|p\.a\.?|annually)',
    ]
    text_lower = text.lower()
    for p in patterns:
        m = re.search(p, text_lower)
        if m:
            val = float(m.group(1).replace(",", ""))
            if val > 1000:
                val = val / 100000  # Convert from absolute to LPA
            return round(val, 2)
    return 0.0

def extract_notice_period(text: str) -> int:
    """Return notice period in days."""
    text_lower = text.lower()
    if re.search(r'immediate|available\s+immediately|notice\s*:\s*0|no\s+notice', text_lower):
        return 0
    m = re.search(r'notice\s+period\s*[:\-]?\s*(\d+)\s*(days?|months?)', text_lower)
    if m:
        val, unit = int(m.group(1)), m.group(2)
        return val * 30 if "month" in unit else val
    if "30 days" in text_lower or "one month" in text_lower or "1 month" in text_lower:
        return 30
    if "60 days" in text_lower or "two months" in text_lower or "2 months" in text_lower:
        return 60
    if "90 days" in text_lower or "three months" in text_lower or "3 months" in text_lower:
        return 90
    return 60  # Default 

def extract_location(text: str) -> str:
    """Extract city from resume text."""
    CITIES = [
        "hyderabad", "secunderabad", "bengaluru", "bangalore",
        "chennai", "mumbai", "pune", "delhi", "noida", "gurgaon",
        "gurugram", "kolkata", "ahmedabad", "coimbatore", "kochi",
        "thiruvananthapuram", "vizag", "visakhapatnam", "vijayawada",
        "warangal", "tirupati", "nagpur", "indore", "bhopal",
        "jaipur", "lucknow", "chandigarh", "mysuru", "mysore",
    ]
    text_lower = text.lower()
    for city in CITIES:
        if city in text_lower:
            return city.title()
    return ""

# def parse_resume(pdf_file) -> Dict[str, Any]:
#     """Full resume parser — returns structured dict."""
#     if hasattr(pdf_file, "read"):
#         raw_bytes = pdf_file.read()
#         pdf_file.seek(0)
#     else:
#         raw_bytes = pdf_file

#     text = extract_text_from_pdf(io.BytesIO(raw_bytes))
#     if not text.strip():
#         return {"name": "Unknown", "error": "Could not extract text from PDF", "raw_text": ""}

def parse_resume(pdf_file, file_type="pdf") -> Dict[str, Any]:
    """Full resume parser — returns structured dict. Supports PDF and DOCX."""
    if hasattr(pdf_file, "read"):
        raw_bytes = pdf_file.read()
        pdf_file.seek(0)
    else:
        raw_bytes = pdf_file

    if file_type == "docx":
        text = extract_text_from_docx(io.BytesIO(raw_bytes))
    else:
        text = extract_text_from_pdf(io.BytesIO(raw_bytes))

    if not text.strip():
        return {"name": "Unknown", "error": "Could not extract text from file", "raw_text": ""}

    return {
        "name": extract_name(text),
        "email": extract_email(text),
        "phone": extract_phone(text),
        "skills": extract_skills(text),
        "education": extract_education(text),
        "experience_years": extract_experience_years(text),
        "current_ctc": extract_ctc(text, "current"),
        "expected_ctc": extract_ctc(text, "expected"),
        "notice_period": extract_notice_period(text),
        "location": extract_location(text),   
        "raw_text": text[:5000],  # Store first 5000 chars
    }
